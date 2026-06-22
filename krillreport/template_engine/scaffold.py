"""Generate a starter ``.docx`` layout template with anchors pre-placed.

Gives operators a known-good, fully-styled template to customise in Word: a cover with
scalar placeholders, a short legend of the available tokens, and the block anchors that
:class:`~krillreport.report_renderer.docx_template_renderer.DocxTemplateRenderer` fills
in. Editing this file (colours, fonts, cover art, headers/footers, moving/removing
anchors) is the low-effort path to a brand-faithful report.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..report_renderer.docx_template_renderer import BLOCK_TOKENS

# (token, heading shown above the anchor in the scaffold) — order = report order.
_BLOCK_LABELS = {
    "executive_summary": "Executive Summary",
    "scope": "Scope",
    "methodology": "Methodology",
    "findings_summary": "Findings Summary",
    "findings": "Detailed Findings",
    "asset_inventory": "Asset Inventory",
    "conclusion": "Conclusion",
    "appendices": "Appendices",
}

_SCALARS = (
    "report_title", "client", "project", "engagement_type",
    "classification", "date", "date_range", "version", "assessors", "overall_risk",
)


def build_scaffold_template(path: Path) -> Path:
    """Write a starter layout template to ``path`` and return it."""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # --- Cover -------------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("{{report_title}}")
    run.bold = True
    run.font.size = Pt(28)

    sub = doc.add_paragraph("{{engagement_type}} — {{client}}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for label, token in (
        ("Project", "project"), ("Classification", "classification"),
        ("Assessment window", "date_range"), ("Overall risk", "overall_risk"),
        ("Version", "version"), ("Date issued", "date"), ("Assessors", "assessors"),
    ):
        para = doc.add_paragraph()
        para.add_run(f"{label}: ").bold = True
        para.add_run("{{" + token + "}}")

    doc.add_page_break()

    # --- How-to legend (delete in a real template) -------------------------
    doc.add_heading("How to use this template", level=1)
    doc.add_paragraph(
        "Style this document however you like in Word — fonts, colours, cover art, "
        "headers/footers. KrillReport keeps all of it and only fills in the tokens below. "
        "Each block anchor must sit alone on its own line. Delete this section."
    )
    doc.add_paragraph(
        "Reference a token by wrapping its name in double braces (shown without braces "
        "here so this guide isn't itself filled in). Scalar tokens are replaced inline:"
    )
    for token in _SCALARS:
        doc.add_paragraph(token, style="List Bullet")
    doc.add_paragraph(
        "Block anchors expand to a whole section (heading included, styled by this "
        "template's Heading styles); each must sit alone on its own line:"
    )
    for token in BLOCK_TOKENS:
        doc.add_paragraph(f"{token} — {_BLOCK_LABELS[token]}", style="List Bullet")
    doc.add_page_break()

    # --- Body: one anchor per section, each alone on its own line ----------
    # Each anchor expands to its section (with its own heading); reorder or delete freely.
    for token in BLOCK_TOKENS:
        doc.add_paragraph("{{" + token + "}}")

    doc.save(str(path))
    return path
