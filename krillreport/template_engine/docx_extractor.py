"""Extract branding (fonts, colours, logo, chrome) from a sample ``.docx``.

python-docx exposes styles, runs, headers/footers and embedded image parts, which is
enough to approximate a house style. Colour extraction is heuristic: we tally the
explicit (non-grey, non-black) run/style colours and treat the most frequent as the
primary brand colour. Everything is best-effort and wrapped so a quirky document never
raises — missing values simply fall back to defaults downstream.
"""

from __future__ import annotations

import re
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document

from ..logging_config import get_logger

logger = get_logger(__name__)

# Heading style names to probe, in order of preference.
_HEADING_STYLES = ("Heading 1", "Title", "Heading 2")
# Theme scheme slots that carry a brand colour, in priority order (skip dk1/lt1/lt2,
# which are body text / page background, not accents).
_THEME_SLOTS = ("dk2", "accent1", "accent2", "accent5", "accent6", "accent3", "accent4")


def extract_docx_branding(path: Path, dest_dir: Path) -> Dict[str, Any]:
    """Return a dict of branding fields discovered in ``path`` (logo saved to ``dest_dir``)."""
    branding: Dict[str, Any] = {}
    try:
        doc = Document(str(path))
    except Exception as exc:
        logger.warning("Could not open DOCX template %s: %s", path.name, exc)
        return branding

    _extract_fonts(doc, branding)
    _extract_colors(doc, branding, _read_theme_colors(path), _read_shading_fills(path))
    _extract_chrome(doc, branding)
    _extract_page_size(doc, branding)

    logo = _extract_logo(doc, dest_dir)
    if logo:
        branding["logo_path"] = str(logo)

    logger.info("Extracted DOCX branding from %s: %s", path.name, sorted(branding))
    return branding


# --------------------------------------------------------------------------- #


def _extract_fonts(doc, branding: Dict[str, Any]) -> None:
    try:
        normal = doc.styles["Normal"].font.name
        if normal:
            branding["body_font"] = normal
    except (KeyError, AttributeError):
        pass
    for style_name in _HEADING_STYLES:
        try:
            font_name = doc.styles[style_name].font.name
        except (KeyError, AttributeError):
            continue
        if font_name:
            branding["heading_font"] = font_name
            break


def _extract_colors(
    doc, branding: Dict[str, Any], theme: Dict[str, str], shading: List[str]
) -> None:
    """Pick primary/secondary/accent brand colours, preferring *intentional* sources.

    Tallying every coloured body run is unreliable: pentest templates routinely embed a
    severity legend (red/amber/green swatches) that out-counts the real house colour. So
    we trust deliberate style/theme choices first — explicit heading-style colours, then
    the document theme scheme, then table shading — and only fall back to body-run
    frequency when a document carries no theme or styled headings at all.
    """
    ordered: List[str] = []

    def add(value: Optional[str]) -> None:
        if not value:
            return
        hex_color = _to_hex(value)
        if _is_brand_color(hex_color.lstrip("#")) and hex_color not in ordered:
            ordered.append(hex_color)

    # 1. Colours explicitly set on heading styles — the strongest signal of intent.
    for style_name in _HEADING_STYLES:
        try:
            color = doc.styles[style_name].font.color
            if color is not None and color.rgb is not None:
                add(str(color.rgb))
        except (KeyError, AttributeError, ValueError):
            continue

    # 2. The document theme's brand slots (dk2 / accentN).
    for slot in _THEME_SLOTS:
        add(theme.get(slot))

    # 3. Table-shading fills (header bands etc.) as a weaker hint.
    for fill in shading:
        add(fill)

    # 4. Last resort only: the most frequent explicit body-run colour.
    if not ordered:
        counter: Counter = Counter()
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                hex_color = _run_color_hex(run)
                if hex_color:
                    counter[hex_color] += 1
        for color, _ in counter.most_common():
            add(color)

    for field, value in zip(("primary_color", "secondary_color", "accent_color"), ordered):
        branding[field] = value


def _read_theme_colors(path: Path) -> Dict[str, str]:
    """Read the ``<a:clrScheme>`` brand colours from the package's theme part."""
    colors: Dict[str, str] = {}
    try:
        with zipfile.ZipFile(str(path)) as zf:
            themes = [n for n in zf.namelist() if n.startswith("word/theme/") and n.endswith(".xml")]
            if not themes:
                return colors
            xml = zf.read(sorted(themes)[0]).decode("utf-8", "replace")
    except (OSError, KeyError, zipfile.BadZipFile):
        return colors
    scheme = re.search(r"<a:clrScheme.*?</a:clrScheme>", xml, re.S)
    if not scheme:
        return colors
    for slot, body in re.findall(
        r"<a:(dk1|lt1|dk2|lt2|accent[1-6]|hlink|folHlink)>(.*?)</a:\1>", scheme.group(0), re.S
    ):
        match = re.search(r'(?:srgbClr val="|lastClr=")([0-9A-Fa-f]{6})', body)
        if match:
            colors[slot] = match.group(1).upper()
    return colors


def _read_shading_fills(path: Path) -> List[str]:
    """Return distinct non-white table/paragraph shading fills, most frequent first."""
    try:
        with zipfile.ZipFile(str(path)) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", "replace")
    except (OSError, KeyError, zipfile.BadZipFile):
        return []
    counter = Counter(re.findall(r'<w:shd[^>]*w:fill="([0-9A-Fa-f]{6})"', xml))
    return [c.upper() for c, _ in counter.most_common()]


def _run_color_hex(run) -> Optional[str]:
    try:
        color = run.font.color
        if color is None or color.rgb is None:
            return None
        return str(color.rgb).upper()
    except (AttributeError, ValueError):
        # THEME colours raise on .rgb access in some python-docx versions.
        return None


def _extract_chrome(doc, branding: Dict[str, Any]) -> None:
    if not doc.sections:
        return
    section = doc.sections[0]
    try:
        header_text = " ".join(p.text for p in section.header.paragraphs).strip()
        if header_text:
            branding["header_text"] = header_text
    except Exception:  # pragma: no cover - defensive
        pass
    try:
        footer_text = " ".join(p.text for p in section.footer.paragraphs).strip()
        if footer_text:
            branding["footer_text"] = footer_text
    except Exception:  # pragma: no cover - defensive
        pass


def _extract_page_size(doc, branding: Dict[str, Any]) -> None:
    if not doc.sections:
        return
    try:
        width_mm = doc.sections[0].page_width.mm
    except (AttributeError, TypeError):
        return
    if width_mm is None:
        return
    # US Letter is 215.9mm wide; A4 is 210mm. Use a midpoint threshold.
    branding["page_size"] = "letter" if width_mm > 213 else "A4"


def _extract_logo(doc, dest_dir: Path) -> Optional[Path]:
    """Save the first embedded image (likely the logo) and return its path."""
    for content_type, blob in _iter_image_blobs(doc):
        ext = _ext_for_content_type(content_type)
        if not ext or not blob:
            continue
        dest_dir.mkdir(parents=True, exist_ok=True)
        logo_path = dest_dir / f"logo{ext}"
        try:
            logo_path.write_bytes(blob)
            return logo_path
        except OSError as exc:  # pragma: no cover
            logger.debug("Failed to write extracted logo: %s", exc)
    return None


def _iter_image_blobs(doc):
    """Yield ``(content_type, blob)`` for image parts in the document and its headers."""
    parts: List[Any] = [doc.part]
    for section in doc.sections:
        for chrome in (section.header, section.footer):
            try:
                parts.append(chrome.part)
            except Exception:  # pragma: no cover
                continue
    seen = set()
    for part in parts:
        related = getattr(part, "related_parts", {})
        for related_part in related.values():
            content_type = getattr(related_part, "content_type", "")
            if not content_type.startswith("image/"):
                continue
            blob = getattr(related_part, "blob", None)
            key = id(related_part)
            if key in seen:
                continue
            seen.add(key)
            yield content_type, blob


def _ext_for_content_type(content_type: str) -> Optional[str]:
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/x-emf": None,  # EMF/WMF aren't usable by the renderers; skip.
        "image/x-wmf": None,
    }
    return mapping.get(content_type, ".png" if content_type.startswith("image/") else None)


# --------------------------------------------------------------------------- #
# Colour helpers
# --------------------------------------------------------------------------- #


def _to_hex(value: str) -> str:
    return "#" + value.lstrip("#").upper()


def _is_brand_color(hex_value: str) -> bool:
    """Exclude near-black, near-white and grey colours — they are body text, not brand."""
    r, g, b = _hex_to_rgb(hex_value)
    if max(r, g, b) - min(r, g, b) < 24:  # grey / mono
        return False
    if r + g + b < 80:  # near-black
        return False
    if min(r, g, b) > 232:  # near-white
        return False
    return True


def _hex_to_rgb(hex_value: str) -> Tuple[int, int, int]:
    text = hex_value.lstrip("#")
    if len(text) != 6:
        return (0, 0, 0)
    try:
        return tuple(int(text[i : i + 2], 16) for i in (0, 2, 4))  # type: ignore[return-value]
    except ValueError:
        return (0, 0, 0)
