"""Render a report *into* a customer's ``.docx`` template (layout fidelity).

Unlike :class:`DocxRenderer`, which builds a document from scratch with KrillReport's own
layout, this renderer opens the customer's template as the base document — inheriting its
theme, styles, fonts, headers/footers, and any cover/boilerplate — and fills it in:

* **Scalar tokens** (``{{report_title}}``, ``{{client}}``, ``{{date}}`` …) are replaced
  inline wherever they appear (cover, headers, footers).
* **Block anchors** (``{{executive_summary}}``, ``{{findings}}``, ``{{asset_inventory}}``
  …) are replaced by the generated section, built with the template's own styles and
  inserted exactly where the anchor sits.

The section *content* is produced by reusing :class:`DocxRenderer`'s builders against a
scratch document, then the resulting block elements are deep-copied in after each anchor
(style references resolve against the template). If a template carries no block anchors at
all, every standard section is appended after its existing content — so even an
un-annotated but well-styled template renders a complete, on-brand report.
"""

from __future__ import annotations

import copy
from pathlib import Path
from typing import Callable, List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from ..logging_config import get_logger
from ..models import NormalizedReport
from ..template_engine import Branding, default_branding
from .docx_renderer import DocxRenderer
from .sections import build_context

logger = get_logger(__name__)

# Block-anchor token -> the order it appends in the no-anchor fallback.
BLOCK_TOKENS = (
    "executive_summary",
    "scope",
    "methodology",
    "findings_summary",
    "findings",
    "asset_inventory",
    "conclusion",
    "appendices",
)


class DocxTemplateRenderer:
    """Fill a customer ``.docx`` template with the report's content."""

    def render(
        self,
        report: NormalizedReport,
        branding: Optional[Branding],
        output_path: Path,
        template_path: Path,
    ) -> Path:
        branding = branding or default_branding()
        context = build_context(report, branding)
        base = DocxRenderer()
        builders = self._builders(base, branding, context, report)

        doc = Document(str(template_path))
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        self._replace_scalars(doc, context)

        injected = 0
        for token, build_fn in builders.items():
            anchor = self._find_anchor(doc, token)
            if anchor is None:
                continue
            self._insert_elements(anchor, self._build_elements(build_fn))
            anchor._p.getparent().remove(anchor._p)
            injected += 1

        if injected == 0:
            # No anchors at all → append the full report after the template content.
            for token in BLOCK_TOKENS:
                for element in self._build_elements(builders[token]):
                    doc.element.body.append(element)

        doc.save(str(output_path))
        logger.info(
            "Wrote DOCX from template %s (%d anchor(s) filled): %s",
            template_path.name, injected, output_path,
        )
        return output_path

    # ------------------------------------------------------------------ #

    def _builders(self, base, branding, context, report):
        """Map each block token to a function that appends its content to a doc."""
        md = report.metadata
        return {
            "executive_summary": lambda d: base._narrative(d, "Executive Summary", md.executive_summary, branding),
            "scope": lambda d: base._scope(d, context),
            "methodology": lambda d: base._narrative(d, "Methodology", md.methodology, branding),
            "findings_summary": lambda d: base._findings_summary(d, branding, context) if context["has_findings"] else None,
            "findings": lambda d: base._detailed_findings(d, branding, context) if context["has_findings"] else None,
            "asset_inventory": lambda d: base._asset_inventory(d, branding, context) if context["has_hosts"] else None,
            "conclusion": lambda d: base._narrative(d, "Conclusion", md.conclusion, branding),
            "appendices": lambda d: base._appendices(d, report, branding),
        }

    def _build_elements(self, build_fn: Callable) -> List:
        """Run ``build_fn`` against a scratch doc; return deep-copied block elements."""
        scratch = Document()
        build_fn(scratch)
        elements: List = []
        for child in list(scratch.element.body):
            if child.tag == qn("w:sectPr"):
                continue
            # Drop the scratch document's leading empty paragraph.
            if child.tag == qn("w:p") and not Paragraph(child, scratch).text.strip() \
                    and child.find(qn("w:r")) is None:
                continue
            elements.append(copy.deepcopy(child))
        return elements

    def _insert_elements(self, anchor: Paragraph, elements: List) -> None:
        cursor = anchor._p
        for element in elements:
            cursor.addnext(element)
            cursor = element

    def _find_anchor(self, doc, token: str) -> Optional[Paragraph]:
        needle = "{{" + token + "}}"
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == needle:
                return paragraph
        return None

    def _replace_scalars(self, doc, context) -> None:
        scalars = {
            "report_title": context["title"],
            "client": context["client"],
            "project": context["metadata"].project_name,
            "engagement_type": context["engagement_type"],
            "classification": context["classification"],
            "date": context["generated"],
            "date_range": context["date_range"],
            "version": context["version"],
            "assessors": context["assessors"],
            "overall_risk": context["posture_label"],
        }
        braced = {"{{" + key + "}}": str(value or "") for key, value in scalars.items()}

        paragraphs = list(doc.paragraphs)
        for section in doc.sections:
            for chrome in (
                section.header, section.footer,
                section.first_page_header, section.first_page_footer,
                section.even_page_header, section.even_page_footer,
            ):
                paragraphs.extend(chrome.paragraphs)

        for paragraph in paragraphs:
            text = paragraph.text
            if "{{" not in text:
                continue
            stripped = text.strip()
            if stripped in braced:
                # A standalone placeholder line — replace it whole, keeping its style.
                _set_paragraph_text(paragraph, braced[stripped])
                continue
            # Inline token within mixed content — replace per run (keeps fields/formatting).
            for run in paragraph.runs:
                if "{{" in run.text:
                    new_text = run.text
                    for token, value in braced.items():
                        new_text = new_text.replace(token, value)
                    run.text = new_text


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace a paragraph's text, preserving its first run's formatting + the style."""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for extra in runs[1:]:
            extra._r.getparent().remove(extra._r)
    else:
        paragraph.add_run(text)
