"""Render a :class:`NormalizedReport` to a branded ``.docx`` via python-docx.

python-docx has no high-level concept of cover pages, coloured table cells, TOC fields
or page numbers, so this module includes the small OOXML helpers needed to produce a
polished document: cell shading, ``PAGE``/``TOC`` fields, and styled headings driven by
the selected :class:`Branding`.
"""

from __future__ import annotations

from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ..logging_config import get_logger
from ..models import Finding, NormalizedReport
from ..template_engine import Branding, default_branding
from ..template_engine.branding import hex_to_rgb
from .markdown_render import inline_segments, parse_blocks
from .sections import build_context

logger = get_logger(__name__)


# --------------------------------------------------------------------------- #
# OOXML helpers
# --------------------------------------------------------------------------- #


def _rgb(hex_color: str) -> RGBColor:
    r, g, b = hex_to_rgb(hex_color)
    return RGBColor(r, g, b)


def _set_cell_background(cell, hex_color: str) -> None:
    """Apply a solid fill colour to a table cell."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def _shade_paragraph(paragraph, hex_color: str) -> None:
    """Apply a background shading colour to a paragraph (used for evidence blocks)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    paragraph.paragraph_format.element.get_or_add_pPr().append(shd)


def _apply_table_style(table, preferred: str) -> None:
    """Apply a table style, falling back gracefully if the named style is absent.

    Not every built-in Word table style is guaranteed to exist in python-docx's
    default template, so we try the preferred style, then progressively plainer ones.
    """
    for name in (preferred, "Light Grid", "Table Grid"):
        try:
            table.style = name
            return
        except KeyError:
            continue


def _add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    """Append a Word field (e.g. ``PAGE``, ``NUMPAGES``, ``TOC``) to a paragraph."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


# --------------------------------------------------------------------------- #
# Renderer
# --------------------------------------------------------------------------- #


class DocxRenderer:
    """Render reports to DOCX."""

    def render(
        self,
        report: NormalizedReport,
        branding: Optional[Branding],
        output_path: Path,
    ) -> Path:
        branding = branding or default_branding()
        context = build_context(report, branding)
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        self._configure_styles(doc, branding)
        self._configure_page_chrome(doc, branding, context)

        self._cover(doc, branding, context)
        self._toc(doc)

        self._narrative(doc, "Executive Summary", report.metadata.executive_summary, branding)
        self._scope(doc, context)
        self._narrative(doc, "Methodology", report.metadata.methodology, branding)

        if context["has_findings"]:
            self._findings_summary(doc, branding, context)
            self._detailed_findings(doc, branding, context)
        else:
            doc.add_heading("Findings Summary", level=1)
            doc.add_paragraph("No findings were identified or imported for this engagement.")

        if context["has_hosts"]:
            self._asset_inventory(doc, branding, context)

        self._narrative(doc, "Conclusion", report.metadata.conclusion, branding)
        self._appendices(doc, report, branding)

        doc.save(str(output_path))
        logger.info("Wrote DOCX report: %s", output_path)
        return output_path

    # ------------------------------------------------------------------ #
    # Setup
    # ------------------------------------------------------------------ #

    def _configure_styles(self, doc: Document, branding: Branding) -> None:
        normal = doc.styles["Normal"]
        normal.font.name = branding.body_font
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = _rgb(branding.text_color)

        for style_name, color in (
            ("Heading 1", branding.primary_color),
            ("Heading 2", branding.primary_color),
            ("Heading 3", branding.secondary_color),
        ):
            try:
                style = doc.styles[style_name]
            except KeyError:
                continue
            style.font.name = branding.heading_font
            style.font.color.rgb = _rgb(color)

        # Letter vs A4 page size.
        section = doc.sections[0]
        if branding.page_size.lower() == "a4":
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
        else:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11.0)

    def _configure_page_chrome(self, doc: Document, branding: Branding, context) -> None:
        section = doc.sections[0]
        header_text = branding.header_text or context["classification"]
        footer_text = branding.footer_text or f"{context['classification']} — {context['title']}"

        header_para = section.header.paragraphs[0]
        header_para.text = header_text
        for run in header_para.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = _rgb(branding.muted_color)

        footer_para = section.footer.paragraphs[0]
        footer_para.text = ""
        right_tab = section.page_width - section.left_margin - section.right_margin
        footer_para.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)
        run = footer_para.add_run(footer_text + "\t")
        run.font.size = Pt(8)
        run.font.color.rgb = _rgb(branding.muted_color)
        page_run = footer_para.add_run("Page ")
        page_run.font.size = Pt(8)
        page_run.font.color.rgb = _rgb(branding.muted_color)
        _add_field(footer_para, "PAGE")
        of_run = footer_para.add_run(" of ")
        of_run.font.size = Pt(8)
        of_run.font.color.rgb = _rgb(branding.muted_color)
        _add_field(footer_para, "NUMPAGES")

    # ------------------------------------------------------------------ #
    # Sections
    # ------------------------------------------------------------------ #

    def _cover(self, doc: Document, branding: Branding, context) -> None:
        if branding.logo_path and Path(branding.logo_path).exists():
            try:
                doc.add_picture(branding.logo_path, width=Inches(2.0))
            except Exception as exc:  # pragma: no cover - bad image
                logger.debug("Could not embed cover logo: %s", exc)

        for _ in range(2):
            doc.add_paragraph()

        title_para = doc.add_paragraph()
        title_run = title_para.add_run(context["title"])
        title_run.font.name = branding.heading_font
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = _rgb(branding.primary_color)

        sub_para = doc.add_paragraph()
        subtitle = context["engagement_type"]
        if context["client"]:
            subtitle += f" — {context['client']}"
        sub_run = sub_para.add_run(subtitle)
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = _rgb(branding.secondary_color)

        # Classification call-out.
        cls_para = doc.add_paragraph()
        cls_run = cls_para.add_run(f"  {context['classification']}  ")
        cls_run.font.bold = True
        cls_run.font.size = Pt(10)
        cls_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Shade the run by shading the whole paragraph with the accent colour.
        _shade_paragraph(cls_para, branding.accent_color)

        doc.add_paragraph()

        meta_rows: List[Tuple[str, str]] = [("Engagement Type", context["engagement_type"])]
        if context["client"]:
            meta_rows.insert(0, ("Client", context["client"]))
        if context["metadata"].project_name:
            meta_rows.append(("Project", context["metadata"].project_name))
        meta_rows.extend(
            [
                ("Assessment Window", context["date_range"]),
                ("Overall Risk", context["posture_label"]),
                ("Report Version", context["version"]),
                ("Date Issued", context["generated"]),
            ]
        )
        if context["assessors"]:
            meta_rows.append(("Assessors", context["assessors"]))

        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for key, value in meta_rows:
            row = table.add_row().cells
            k_run = row[0].paragraphs[0].add_run(key)
            k_run.font.bold = True
            k_run.font.color.rgb = _rgb(branding.muted_color)
            k_run.font.size = Pt(9)
            row[1].paragraphs[0].add_run(str(value)).font.size = Pt(10)
            row[0].width = Inches(2.2)
            row[1].width = Inches(4.0)

        doc.add_page_break()

    def _toc(self, doc: Document) -> None:
        doc.add_heading("Contents", level=1)
        para = doc.add_paragraph()
        _add_field(
            para,
            'TOC \\o "1-2" \\h \\z \\u',
            placeholder='Right-click here and choose "Update Field" to build the table of contents.',
        )
        doc.add_page_break()

    def _narrative(self, doc: Document, title: str, text: str, branding: Branding) -> None:
        if not text or not text.strip():
            return
        doc.add_heading(title, level=1)
        self._markdown(doc, text, branding)

    def _scope(self, doc: Document, context) -> None:
        scope = context["scope"]
        if not scope:
            return
        doc.add_heading("Scope", level=1)
        doc.add_paragraph("The following targets were in scope for this engagement:")
        for item in scope:
            doc.add_paragraph(str(item), style="List Bullet")

    def _findings_summary(self, doc: Document, branding: Branding, context) -> None:
        doc.add_heading("Findings Summary", level=1)
        summary = context["summary"]
        doc.add_paragraph(
            f"A total of {summary.total_findings} finding(s) were identified. "
            f"The overall risk posture is {context['posture_label']}."
        )

        # Severity breakdown table with coloured severity cells + a simple bar.
        rows = context["severity_rows"]
        max_count = max((r["count"] for r in rows), default=0) or 1
        table = doc.add_table(rows=1, cols=3)
        _apply_table_style(table, "Light Grid Accent 1")
        hdr = table.rows[0].cells
        for i, label in enumerate(("Severity", "Count", "Distribution")):
            run = hdr[i].paragraphs[0].add_run(label)
            run.font.bold = True
        for row in rows:
            cells = table.add_row().cells
            sev_run = cells[0].paragraphs[0].add_run(row["label"])
            sev_run.font.bold = True
            sev_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_background(cells[0], row["color"])
            cells[1].paragraphs[0].add_run(str(row["count"]))
            bar_len = round(row["count"] / max_count * 20) if row["count"] else 0
            bar_run = cells[2].paragraphs[0].add_run("█" * bar_len)
            bar_run.font.color.rgb = _rgb(row["color"])

        doc.add_paragraph()

        # Overview table of all findings.
        overview = doc.add_table(rows=1, cols=5)
        _apply_table_style(overview, "Light Grid Accent 1")
        headers = overview.rows[0].cells
        for i, label in enumerate(("#", "Finding", "Severity", "CVSS", "Affected")):
            run = headers[i].paragraphs[0].add_run(label)
            run.font.bold = True
        for num, finding in context["numbered_findings"]:
            cells = overview.add_row().cells
            cells[0].paragraphs[0].add_run(str(num))
            cells[1].paragraphs[0].add_run(finding.title)
            sev_run = cells[2].paragraphs[0].add_run(finding.severity.value)
            sev_run.font.bold = True
            sev_run.font.color.rgb = _rgb(branding.severity_color(finding.severity))
            cells[3].paragraphs[0].add_run(
                f"{finding.cvss_score:.1f}" if finding.cvss_score is not None else "—"
            )
            cells[4].paragraphs[0].add_run(
                ", ".join(finding.affected_assets) if finding.affected_assets else "—"
            )

    def _detailed_findings(self, doc: Document, branding: Branding, context) -> None:
        doc.add_heading("Detailed Findings", level=1)
        for num, finding in context["numbered_findings"]:
            scan_rows = context["finding_scan_rows"].get(finding.id)
            self._finding(doc, branding, num, finding, scan_rows)

    def _finding(
        self,
        doc: Document,
        branding: Branding,
        num: int,
        finding: Finding,
        scan_rows: Optional[List[str]] = None,
    ) -> None:
        doc.add_heading(f"{num}. {finding.title}", level=2)

        # Severity call-out line.
        sev_para = doc.add_paragraph()
        sev_run = sev_para.add_run(f"Severity: {finding.severity.value}")
        sev_run.font.bold = True
        sev_run.font.color.rgb = _rgb(branding.severity_color(finding.severity))

        # Metadata table.
        meta_rows: List[Tuple[str, str]] = []
        cvss = f"{finding.cvss_score:.1f}" if finding.cvss_score is not None else "—"
        if finding.cvss_vector:
            cvss += f" ({finding.cvss_vector})"
        meta_rows.append(("CVSS", cvss))
        if finding.category:
            meta_rows.append(("Category", finding.category))
        meta_rows.append(("Status", finding.status.value))
        if finding.cve:
            meta_rows.append(("CVE", ", ".join(finding.cve)))
        if finding.cwe:
            meta_rows.append(("CWE", ", ".join(finding.cwe)))
        if finding.affected_assets:
            meta_rows.append(("Affected", ", ".join(finding.affected_assets)))
        # Correlated scan data (nmap etc.) for the affected assets.
        if scan_rows:
            meta_rows.append(("Scan Data", "\n".join(scan_rows)))

        table = doc.add_table(rows=0, cols=2)
        _apply_table_style(table, "Light List Accent 1")
        for key, value in meta_rows:
            cells = table.add_row().cells
            k_run = cells[0].paragraphs[0].add_run(key)
            k_run.font.bold = True
            cells[0].width = Inches(1.4)
            cells[1].paragraphs[0].add_run(value)
            cells[1].width = Inches(5.0)

        self._labelled_prose(doc, "Description", finding.description, branding)
        self._labelled_prose(doc, "Impact", finding.impact, branding)

        if finding.evidence:
            doc.add_heading("Evidence", level=3)
            for evidence in finding.evidence:
                if evidence.caption:
                    cap = doc.add_paragraph()
                    cap_run = cap.add_run(evidence.caption)
                    cap_run.font.italic = True
                    cap_run.font.size = Pt(9)
                    cap_run.font.color.rgb = _rgb(branding.muted_color)
                if evidence.text:
                    para = doc.add_paragraph()
                    _shade_paragraph(para, "#1E2530")
                    ev_run = para.add_run(evidence.text)
                    ev_run.font.name = branding.mono_font
                    ev_run.font.size = Pt(8.5)
                    ev_run.font.color.rgb = RGBColor(0xE6, 0xED, 0xF3)
                if evidence.image_path and Path(evidence.image_path).exists():
                    try:
                        doc.add_picture(evidence.image_path, width=Inches(6.0))
                    except Exception as exc:  # pragma: no cover
                        logger.debug("Could not embed evidence image: %s", exc)

        self._labelled_prose(doc, "Remediation", finding.remediation, branding)

        if finding.references:
            doc.add_heading("References", level=3)
            for ref in finding.references:
                text = ref.title or ref.url
                if ref.url and ref.url != text:
                    text = f"{text} — {ref.url}"
                doc.add_paragraph(text, style="List Bullet")

    def _labelled_prose(self, doc: Document, label: str, text: str, branding: Branding) -> None:
        if not text or not text.strip():
            return
        doc.add_heading(label, level=3)
        self._markdown(doc, text, branding)

    def _markdown(self, doc: Document, text: str, branding: Branding) -> None:
        """Emit Markdown prose as native DOCX blocks (paragraphs, lists, tables, code)."""
        for block in parse_blocks(text):
            kind = block["kind"]
            if kind == "heading":
                para = doc.add_paragraph()
                _add_inline_runs(para, block["text"], branding)
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = _rgb(branding.secondary_color)
            elif kind == "para":
                para = doc.add_paragraph()
                _add_inline_runs(para, block["text"], branding)
            elif kind in ("ulist", "olist"):
                style = "List Number" if kind == "olist" else "List Bullet"
                for item in block["items"]:
                    para = doc.add_paragraph(style=style)
                    _add_inline_runs(para, item, branding)
            elif kind == "code":
                para = doc.add_paragraph()
                _shade_paragraph(para, "#1E2530")
                run = para.add_run(block["text"])
                run.font.name = branding.mono_font
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0xE6, 0xED, 0xF3)
            elif kind == "quote":
                para = doc.add_paragraph()
                _add_inline_runs(para, block["text"], branding)
                for run in para.runs:
                    run.font.italic = True
                    run.font.color.rgb = _rgb(branding.muted_color)
            elif kind == "table":
                self._markdown_table(doc, block, branding)

    def _markdown_table(self, doc: Document, block, branding: Branding) -> None:
        header = block["header"]
        table = doc.add_table(rows=1, cols=len(header))
        _apply_table_style(table, "Light Grid Accent 1")
        for cell, label in zip(table.rows[0].cells, header):
            _add_inline_runs(cell.paragraphs[0], label, branding)
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
        for row in block["rows"]:
            cells = table.add_row().cells
            for cell, value in zip(cells, row):
                _add_inline_runs(cell.paragraphs[0], value, branding)

    def _asset_inventory(self, doc: Document, branding: Branding, context) -> None:
        doc.add_heading("Asset Inventory", level=1)
        # A "Findings" column cross-references scanned hosts to the findings that hit them.
        show_findings = context["has_findings"]
        labels = ["Host", "IP Address", "Operating System", "Services"]
        if show_findings:
            labels.append("Findings")
        table = doc.add_table(rows=1, cols=len(labels))
        _apply_table_style(table, "Light Grid Accent 1")
        headers = table.rows[0].cells
        for i, label in enumerate(labels):
            run = headers[i].paragraphs[0].add_run(label)
            run.font.bold = True
        host_findings = context["host_findings"]
        for host in context["hosts"]:
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(host.hostname or "—")
            cells[1].paragraphs[0].add_run(host.ip_address or "—")
            cells[2].paragraphs[0].add_run(host.operating_system or "—")
            services = "\n".join(s.label() for s in host.services) or "—"
            cells[3].paragraphs[0].add_run(services)
            if show_findings:
                nums = host_findings.get(host.identifier)
                cells[4].paragraphs[0].add_run(", ".join(str(n) for n in nums) if nums else "—")

    def _appendices(self, doc: Document, report: NormalizedReport, branding: Branding) -> None:
        for index, appendix in enumerate(report.appendices, start=1):
            doc.add_heading(f"Appendix {index}: {appendix.title}", level=1)
            if appendix.language:
                # Verbatim attachment (script/config): one monospaced code block.
                para = doc.add_paragraph()
                _shade_paragraph(para, "#1E2530")
                run = para.add_run(appendix.content)
                run.font.name = branding.mono_font
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0xE6, 0xED, 0xF3)
            else:
                self._markdown(doc, appendix.content, branding)


def _add_inline_runs(paragraph, text: str, branding: Branding) -> None:
    """Add styled runs to ``paragraph`` from inline Markdown (bold/italic/code/links)."""
    for seg in inline_segments(text):
        # A segment may carry hard line breaks; split so each becomes a Word break.
        for line_no, piece in enumerate(seg.text.split("\n")):
            if line_no:
                paragraph.add_run().add_break()
            if not piece:
                continue
            run = paragraph.add_run(piece)
            run.font.bold = seg.bold
            run.font.italic = seg.italic
            if seg.code:
                run.font.name = branding.mono_font
            if seg.href:
                run.font.underline = True
                run.font.color.rgb = _rgb(branding.accent_color)
