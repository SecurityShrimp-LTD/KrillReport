"""Tests for layout-template rendering (scaffold + anchor injection + fallback)."""

from pathlib import Path

from docx import Document

from krillreport.pipeline import run_pipeline
from krillreport.template_engine.scaffold import build_scaffold_template

SAMPLE = Path(__file__).resolve().parents[1] / "examples" / "sample_inputs" / "manual_findings.md"


def _all_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_scaffold_writes_anchors(tmp_path):
    tpl = build_scaffold_template(tmp_path / "layout.docx")
    doc = Document(str(tpl))
    standalone = {p.text.strip() for p in doc.paragraphs}
    assert "{{report_title}}" in standalone  # a scalar placeholder on the cover
    assert "{{findings}}" in standalone       # a block anchor on its own line


def test_anchor_mode_fills_scalars_and_blocks(tmp_path):
    # A minimal template: boilerplate + one scalar + two block anchors.
    doc = Document()
    doc.add_paragraph("ACME PROPRIETARY — boilerplate top")
    doc.add_paragraph("Prepared for: {{client}}")
    doc.add_paragraph("{{executive_summary}}")
    doc.add_paragraph("{{findings}}")
    doc.add_paragraph("boilerplate bottom")
    tpl = tmp_path / "tpl.docx"
    doc.save(str(tpl))

    out = run_pipeline(
        [SAMPLE], output_dir=tmp_path / "out", formats=["docx"], enhance=False,
        metadata_overrides={"client_name": "ACME Corp"}, layout_template=tpl,
    ).outputs["docx"]

    text = _all_text(out)
    assert "{{" not in text                       # every token consumed
    assert "boilerplate top" in text              # template content preserved
    assert "boilerplate bottom" in text
    assert "Prepared for: ACME Corp" in text      # scalar replaced
    assert "Detailed Findings" in text            # block anchor expanded (with heading)
    assert "Invoice" in text                       # a real finding landed


def test_no_anchor_template_appends_full_report(tmp_path):
    doc = Document()
    doc.add_paragraph("Company cover boilerplate")
    tpl = tmp_path / "plain.docx"
    doc.save(str(tpl))

    out = run_pipeline(
        [SAMPLE], output_dir=tmp_path / "out", formats=["docx"], enhance=False,
        layout_template=tpl,
    ).outputs["docx"]

    text = _all_text(out)
    assert "Company cover boilerplate" in text     # template kept
    assert "Detailed Findings" in text             # report appended after it
    assert "Executive Summary" in text


def test_block_order_in_no_anchor_fallback(tmp_path):
    tpl = tmp_path / "p.docx"
    Document().save(str(tpl))
    out = run_pipeline(
        [SAMPLE], output_dir=tmp_path / "o", formats=["docx"], enhance=False, layout_template=tpl,
    ).outputs["docx"]
    headings = [p.text.strip() for p in Document(str(out)).paragraphs]
    # Findings Summary precedes Detailed Findings in the appended order.
    assert headings.index("Findings Summary") < headings.index("Detailed Findings")
