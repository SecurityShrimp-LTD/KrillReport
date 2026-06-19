"""Shared pytest fixtures."""

from __future__ import annotations

from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parents[1]
SAMPLE_INPUTS = REPO_ROOT / "examples" / "sample_inputs"


@pytest.fixture
def sample_inputs() -> list:
    """All shipped example input files."""
    return sorted(p for p in SAMPLE_INPUTS.iterdir() if p.is_file())


@pytest.fixture
def branded_docx(tmp_path: Path) -> Path:
    """Create a minimal branded .docx in tmp_path for branding-extraction tests."""
    from docx import Document
    from docx.shared import RGBColor

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    heading = doc.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading.font.color.rgb = RGBColor(0x0E, 0x7C, 0x7B)  # teal
    doc.sections[0].header.paragraphs[0].text = "ACME — CONFIDENTIAL"
    doc.add_heading("Sample", level=1)
    path = tmp_path / "brand.docx"
    doc.save(str(path))
    return path
