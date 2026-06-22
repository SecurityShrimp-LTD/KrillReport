"""DOCX branding extraction — colour-priority regression tests.

The key risk these guard against: a pentest template usually embeds a severity legend
(red/amber/green swatches) that, by raw frequency, would otherwise be mistaken for the
brand colour. Intentional sources (styled heading colours, the theme scheme) must win.
"""

from pathlib import Path

from docx import Document
from docx.shared import RGBColor

from krillreport.template_engine.docx_extractor import (
    _emf_wmf_ext,
    _ext_for_content_type,
    _is_emf_wmf,
    _read_theme_colors,
    extract_docx_branding,
)


def _legend_doc(tmp_path: Path) -> Path:
    """A document whose most frequent body colours are a red/green/amber legend."""
    doc = Document()
    para = doc.add_paragraph()
    for color in ([RGBColor(0xFF, 0x00, 0x00)] * 12 + [RGBColor(0x00, 0xB0, 0x50)] * 8
                  + [RGBColor(0xFF, 0xC0, 0x00)] * 6):
        run = para.add_run("SEV ")
        run.font.color.rgb = color
    path = tmp_path / "legend.docx"
    doc.save(str(path))
    return path


def test_explicit_heading_colour_beats_body_legend(tmp_path):
    doc = Document()
    doc.styles["Heading 1"].font.color.rgb = RGBColor(0x0E, 0x7C, 0x7B)  # deliberate teal
    para = doc.add_paragraph()
    for _ in range(15):  # a louder red legend than the heading appears
        run = para.add_run("CRITICAL ")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    path = tmp_path / "branded.docx"
    doc.save(str(path))

    branding = extract_docx_branding(path, tmp_path / "assets")
    assert branding["primary_color"].upper() == "#0E7C7B"


def test_severity_legend_never_becomes_primary(tmp_path):
    path = _legend_doc(tmp_path)
    branding = extract_docx_branding(path, tmp_path / "assets")
    primary = branding.get("primary_color", "").upper()
    # The brand colour comes from the heading style / theme, not the legend swatches.
    assert primary not in {"#FF0000", "#00B050", "#FFC000"}
    assert primary  # something sensible was still chosen


def test_read_theme_colors(tmp_path):
    path = _legend_doc(tmp_path)
    theme = _read_theme_colors(path)
    # A real .docx always ships a theme scheme with dk2 + accent slots.
    assert "dk2" in theme and "accent1" in theme
    assert all(len(v) == 6 for v in theme.values())


def test_largest_image_is_chosen_as_logo(tmp_path):
    from docx import Document
    from docx.shared import Inches
    from PIL import Image

    # A tiny icon (few bytes) and a big noisy "logo" (many bytes), embedded in that order.
    icon = tmp_path / "icon.png"
    Image.new("RGB", (16, 16), (200, 30, 30)).save(icon)
    logo = tmp_path / "logo.png"
    Image.effect_noise((600, 300), 80).convert("RGB").save(logo)  # large + incompressible
    assert logo.stat().st_size > icon.stat().st_size

    doc = Document()
    doc.add_picture(str(icon), width=Inches(0.3))   # first in document order
    doc.add_picture(str(logo), width=Inches(3.0))
    src = tmp_path / "doc.docx"
    doc.save(str(src))

    branding = extract_docx_branding(src, tmp_path / "assets")
    chosen = Path(branding["logo_path"])
    # The larger image wins despite the icon coming first.
    assert chosen.stat().st_size == logo.stat().st_size


def test_emf_wmf_classification():
    assert _is_emf_wmf("image/x-emf") and _is_emf_wmf("image/wmf")
    assert not _is_emf_wmf("image/png")
    assert _ext_for_content_type("image/x-emf") is None      # not directly embeddable
    assert _ext_for_content_type("image/png") == ".png"
    assert _emf_wmf_ext("image/x-emf") == ".emf"
    assert _emf_wmf_ext("image/x-wmf") == ".wmf"
