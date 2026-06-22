"""Phase 2 — template-faithful PDF via LibreOffice (with WeasyPrint fallback)."""

from pathlib import Path

from docx import Document

import krillreport.report_renderer as rr
from krillreport.pipeline import run_pipeline
from krillreport.report_renderer.docx_to_pdf import libreoffice_available

SAMPLE = Path(__file__).resolve().parents[1] / "examples" / "sample_inputs" / "manual_findings.md"


def _template(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("COVER {{client}}")
    doc.add_paragraph("{{findings}}")
    path = tmp_path / "tpl.docx"
    doc.save(str(path))
    return path


def test_libreoffice_available_returns_bool():
    assert isinstance(libreoffice_available(), bool)


def test_template_pdf_falls_back_to_weasyprint(tmp_path, monkeypatch):
    monkeypatch.setattr(rr, "libreoffice_available", lambda: False)
    out = run_pipeline(
        [SAMPLE], output_dir=tmp_path / "out", formats=["pdf", "docx"], enhance=False,
        metadata_overrides={"client_name": "ACME"}, layout_template=_template(tmp_path),
    ).outputs
    assert out["pdf"].read_bytes()[:4] == b"%PDF"            # a real PDF was produced
    assert "{{" not in "\n".join(p.text for p in Document(str(out["docx"])).paragraphs)


def test_template_pdf_uses_libreoffice_when_available(tmp_path, monkeypatch):
    calls = {}

    def fake_convert(docx_path, pdf_path):
        calls["docx"] = Path(docx_path)
        Path(pdf_path).write_bytes(b"%PDF-stub-from-libreoffice")
        return Path(pdf_path)

    monkeypatch.setattr(rr, "libreoffice_available", lambda: True)
    monkeypatch.setattr(rr, "convert_docx_to_pdf", fake_convert)

    out = run_pipeline(
        [SAMPLE], output_dir=tmp_path / "out", formats=["pdf"], enhance=False,
        layout_template=_template(tmp_path),
    ).outputs
    # The PDF came from the (stubbed) DOCX→PDF conversion, not WeasyPrint.
    assert out["pdf"].read_bytes() == b"%PDF-stub-from-libreoffice"
    assert calls["docx"].suffix == ".docx"


def test_pdf_only_layout_cleans_intermediate_docx(tmp_path, monkeypatch):
    monkeypatch.setattr(rr, "libreoffice_available", lambda: False)
    out_dir = tmp_path / "out"
    run_pipeline(
        [SAMPLE], output_dir=out_dir, formats=["pdf"], enhance=False,
        layout_template=_template(tmp_path),
    )
    assert list(out_dir.glob("*.pdf"))            # PDF kept
    assert not list(out_dir.glob("*.docx"))       # intermediate DOCX removed
