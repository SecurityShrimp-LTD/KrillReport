#!/usr/bin/env python3
"""Generate a sample branded .docx that KrillReport can extract a template from.

This stands in for the kind of branded report a customer would upload. Run it to
(re)create ``examples/sample_templates/acme_brand.docx``::

    python examples/make_sample_template.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw

HERE = Path(__file__).resolve().parent
OUT_DIR = HERE / "sample_templates"
LOGO = OUT_DIR / "acme_logo.png"
DOCX = OUT_DIR / "acme_brand.docx"

BRAND_TEAL = (0x0E, 0x7C, 0x7B)
BRAND_AMBER = (0xC9, 0x7A, 0x12)


def make_logo() -> None:
    img = Image.new("RGB", (260, 90), "#0E7C7B")
    draw = ImageDraw.Draw(img)
    draw.rectangle([0, 74, 260, 90], fill="#C97A12")
    draw.text((16, 30), "ACME CORP", fill="white")
    img.save(LOGO)


def make_docx() -> None:
    doc = Document()

    # Body + heading fonts.
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    heading = doc.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading.font.color.rgb = RGBColor(*BRAND_TEAL)

    # Header: logo + confidential banner.
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "ACME Corporation — CONFIDENTIAL"
    header.add_paragraph().add_run().add_picture(str(LOGO), width=Inches(1.4))

    # Footer.
    section.footer.paragraphs[0].text = "© ACME Corporation — Page"

    # Body with brand colours so the extractor has colour signal.
    doc.add_heading("Security Assessment Report", level=1)
    para = doc.add_paragraph()
    accent = para.add_run("Confidential — prepared exclusively for ACME Corporation.")
    accent.font.color.rgb = RGBColor(*BRAND_AMBER)
    doc.add_paragraph("This document template demonstrates ACME's house style.")

    doc.save(DOCX)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_logo()
    make_docx()
    print(f"Wrote {LOGO}")
    print(f"Wrote {DOCX}")


if __name__ == "__main__":
    main()
